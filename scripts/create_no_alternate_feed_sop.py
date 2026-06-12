import docx

d = docx.Document()

d.add_heading("No Alternate Feed Available SOP", 0)

d.add_heading("Purpose", level=1)
d.add_paragraph(
    "This SOP provides step-by-step instructions for ground officers when no alternate feed "
    "is available at the tail-end valve. It covers resident notification, affected valve "
    "identification, and temporary water supply arrangements."
)

d.add_heading("Procedure", level=1)

steps = [
    (
        "Step 1 — Notify the operator of no alternate feed",
        "Inform the operator that there is no alternate feed available for this shutdown. The affected downstream area will lose water supply for the duration of the operation.",
    ),
    (
        "Step 2 — Identify affected residents",
        "Identify all residents that draw water supply from the valves downstream of the initial pipe ID. These are all valves in the downstream chain starting from the pipe being shut down.",
    ),
    (
        "Step 3 — List affected valves and road names",
        "Compile and provide the operator with a list of each affected valve ID and its corresponding road name, so that residents along those roads can be notified.",
    ),
    (
        "Step 4 — Arrange temporary water supply",
        "Coordinate the provision of temporary water supply to affected residents via water wagon and water bags for the duration of the shutdown.",
    ),
]

for heading, body in steps:
    d.add_heading(heading, level=2)
    d.add_paragraph(body)

d.save("data/seed/sop_documents/no_alternate_feed_available_SOP.docx")
print("Done.")
