import docx

d = docx.Document()

d.add_heading("Alternate Feed Available SOP", 0)

d.add_heading("Purpose", level=1)
d.add_paragraph(
    "This SOP provides step-by-step valve operation instructions for ground officers "
    "when an alternate feed has been confirmed available at the tail-end valve. "
    "It covers the sequential opening of reverse isolation pipes, working backwards "
    "from the tail-end valve, followed by shutting the valve immediately downstream "
    "of the originally isolated pipe."
)

d.add_heading("Procedure", level=1)

steps = [
    (
        "Step 1 — Identify the first reverse isolation pipe from the tail-end valve",
        "Starting from the tail-end valve, work backwards to identify the first reverse isolation pipe ID. "
        "These are the reverse-direction pipes that were verified as closed during the isolation check.",
    ),
    (
        "Step 2 — Verify the pipe status is closed",
        'Check that the identified pipe status is "closed" before proceeding. '
        "Do not operate the valve if the pipe is already open.",
    ),
    (
        "Step 3 — Operate the downstream valve",
        "Operate the downstream valve of the identified pipe. Confirm the valve has been fully operated.",
    ),
    (
        "Step 4 — Move to the next reverse isolation pipe and repeat",
        "Identify the next reverse isolation pipe working backwards. "
        "Repeat Steps 2 and 3 for each pipe until all reverse isolation pipes have been operated.",
    ),
    (
        "Step 5 — Operate the valve immediately after the originally isolated pipe",
        "Once all reverse isolation pipes have been addressed, operate the valve that sits immediately "
        "downstream of the originally isolated pipe.",
    ),
    (
        "Step 6 — Record the number of valves operated",
        "Record the total number of valves operated during this procedure.",
    ),
]

for heading, body in steps:
    d.add_heading(heading, level=2)
    d.add_paragraph(body)

d.save("data/seed/sop_documents/alternate_feed_available_SOP.docx")
print("Done.")
